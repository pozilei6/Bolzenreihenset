from pathlib import Path
from copy import copy
import json
import re
import sys

from docx import Document
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from openpyxl.utils.cell import range_boundaries


DOCX_FOLDER = Path("Daten/docxDaten")
XLSX_FILE = Path("Daten/xlsxDaten/Grundlagen Bolzenreidhenset.xlsx")
UNMATCHED_FILE = Path("nicht_gematchte_referenzen.txt")
CALC_REPORT_FILE = Path("nicht_geschriebene_berechnungscode_eintraege.txt")
DESCRIPTION_CACHE_FILE = Path("Daten/xlsxDaten/bolzenreihenset_beschreibungen_cache.json")
ABFUELLNORM_SHEET_NAME = "Bolzenreihenset - Abfüllnorm"
CALC_SHEET_NAME = "Bolzenreihenset Berechnungscode"
DRAWING_LOOKUP_SHEET_NAME = "Bolzenartikel - Zeichnungs"


def clean_value(value):
	return value.strip(" \t\r\n\x00\x07")


def read_docx_text(path):
	document = Document(path)
	parts = []

	for paragraph in document.paragraphs:
		parts.append(paragraph.text)

	for table in document.tables:
		for row in table.rows:
			for cell in row.cells:
				parts.append(cell.text)

	return "\x07".join(parts)


def read_legacy_doc_text(path):
	return path.read_bytes().decode("cp1252", errors="ignore")


def read_word_text(path):
	suffix = path.suffix.lower()

	if suffix == ".docx":
		return read_docx_text(path)

	if suffix == ".doc":
		return read_legacy_doc_text(path)

	raise ValueError(f"Unsupported file type: {path}")


def first_match(pattern, text):
	match = re.search(pattern, text, flags=re.IGNORECASE | re.DOTALL)
	return clean_value(match.group(1)) if match else None


def reference_after_funktion(text):
	match = re.search(
		r"Funktion:\s*[^\r\x07\n]+(.{0,200})",
		text,
		flags=re.IGNORECASE | re.DOTALL,
	)
	if not match:
		return None

	reference = re.search(r"\b800\.\d{3}(?:\.\d{3})?(?:\.[A-Za-z])?\b", match.group(1))
	return reference.group(0) if reference else None


def line_from_filename(filename):
	match = re.search(r"(?:^|[_\-\s])L([0-9])(?:[_\-.\s]|$)", filename, flags=re.IGNORECASE)
	return match.group(1) if match else None


def system_from_filename(filename):
	match = re.search(r"(?:^|[_\-\s])K(\d{2,4})(RS|S)?(?=FP\s*2|[_\-.\s]|$)", filename, flags=re.IGNORECASE)
	if not match:
		return None

	return f"{match.group(1)}{match.group(2) or ''}".upper()


def rotor_from_filename(filename):
	match = re.search(r"(?:^|[_\-\s])R(\d+(?:[.,]\d+)?)(ze|ex)?(?=[_\-.\s]|$)", filename, flags=re.IGNORECASE)
	if not match:
		return None

	diameter = match.group(1).replace(",", ".")
	rotor_type = match.group(2)
	return f"{diameter} {rotor_type.lower()}" if rotor_type else diameter


def hulse_from_filename(filename):
	match = re.search(r"(?:^|[_\-\s])H(\d+(?:[.,]\d+)?)(?=[_\-.\s]|$)", filename, flags=re.IGNORECASE)
	return match.group(1).replace(",", ".") if match else None


def can_parse_rotor_value(value):
	try:
		parse_rotor_diameter_and_type(value)
		return True
	except ValueError:
		return False


def can_parse_diameter_value(value, field_name):
	try:
		normalize_diameters(value, field_name)
		return True
	except ValueError:
		return False


def default_line_for_system(system_text):
	try:
		system = normalize_system(system_text)
	except ValueError:
		return None

	if system.startswith("2500"):
		return "1", "2500 Regel"

	if system.startswith("1000"):
		return "1", "1000 Regel"

	return None


def is_keso_casa_word(system_text, filename):
	system = normalize_text(system_text) or ""
	return bool(
		re.search(r"\bKESO\s+CASA\b", system, flags=re.IGNORECASE)
		or re.search(r"KESO[_\-\s]*CASA", filename, flags=re.IGNORECASE)
	)


def is_2500fp_casa_word(system_text, filename):
	system = normalize_text(system_text) or ""
	combined_text = f"{filename} {system}"
	return bool(re.search(r"(?<![A-Za-z0-9])K?2500\s*FP(?![A-Za-z0-9])", combined_text, flags=re.IGNORECASE))


def is_4000s_casa_word(system_text, filename):
	return is_keso_casa_word(system_text, filename) or is_2500fp_casa_word(system_text, filename)


def is_4000s_grundausfuehrung_word(system_text, filename, function_text=None):
	if is_4000s_casa_word(system_text, filename):
		return False

	system = normalize_text(system_text) or ""
	function = normalize_text(function_text) or ""
	combined_text = f"{filename} {system} {function}"
	if re.search(r"Parkeon|Parceon", combined_text, flags=re.IGNORECASE):
		return False

	return bool(
		re.search(r"\b4000S\b", system, flags=re.IGNORECASE)
		or re.search(r"(?:^|[_\-\s])K4000S(?:[_\-\s]|$)", filename, flags=re.IGNORECASE)
	) and bool(re.search(r"Grundausf|Standard", combined_text, flags=re.IGNORECASE))


def word_line_value(text, filename, system_text, function_text):
	line = first_match(r"Linie:\s*([^\r\x07\n]+)", text)
	try:
		normalize_line(line)
		return line, "Dokument"
	except ValueError:
		filename_line = line_from_filename(filename)
		if filename_line:
			return filename_line, "Dateiname"

		if is_4000s_casa_word(system_text, filename):
			return "1", "4000S Casa Regel"

		if is_4000s_grundausfuehrung_word(system_text, filename, function_text):
			return "2", "4000S Grundausfuehrung Regel"

		default_line = default_line_for_system(system_text or system_from_filename(filename))
		if default_line:
			return default_line

		return line, "Dokument"


def parse_word_file(path):
	text = read_word_text(path)
	system = first_match(r"System:\s*([^\r\x07\n]+)", text)
	function = first_match(r"Funktion:\s*([^\r\x07\n]+)", text)
	system = system or system_from_filename(path.name)
	rotor = first_match(r"Rotor[^:]*:\s*([^\r\x07\n]+)", text)
	hulse = first_match(r"H.{0,6}lsen[^:]*:\s*([^\r\x07\n]+)", text)
	filename_rotor = rotor_from_filename(path.name)
	filename_hulse = hulse_from_filename(path.name)
	if filename_rotor and not can_parse_rotor_value(rotor):
		rotor = filename_rotor
	if filename_hulse and not can_parse_diameter_value(hulse, "Hulsen-diam"):
		hulse = filename_hulse
	line, line_source = word_line_value(text, path.name, system, function)

	return {
		"Datei": path.name,
		"System": system,
		"Rotor-diam": rotor,
		"Hulsen-diam": hulse,
		"Linie": line,
		"Linie_Quelle": line_source,
		"Funktion": function,
		"Referenz_zu_Abfullnorm": reference_after_funktion(text),
	}


def load_all_word_data(folder=DOCX_FOLDER):
	word_files = sorted(
		path
		for path in folder.iterdir()
		if path.suffix.lower() in {".doc", ".docx"} and not path.name.startswith("~$")
	)

	return [parse_word_file(path) for path in word_files]


def load_description_cache(path=DESCRIPTION_CACHE_FILE):
	if not path.exists():
		return {}

	return json.loads(path.read_text(encoding="utf-8"))


def load_bolzenreihenset_pairs(path=XLSX_FILE, start_row=2, end_row=None):
	workbook = load_workbook(path, data_only=True, read_only=True)
	sheet = workbook[ABFUELLNORM_SHEET_NAME]
	description_cache = load_description_cache()
	if end_row is None:
		end_row = sheet.max_row

	try:
		return {
			row_number: (
				sheet[f"A{row_number}"].value
				or description_cache.get(str(row_number)),
				sheet[f"B{row_number}"].value,
			)
			for row_number in range(start_row, end_row + 1)
		}
	finally:
		workbook.close()


def normalize_text(value):
	if value is None:
		return None

	text = str(value).replace("\x00", " ")
	return re.sub(r"\s+", " ", clean_value(text))


def normalize_system(value):
	text = normalize_text(value)
	if not text:
		raise ValueError("System fehlt")

	system_match = re.search(
		r"\bK?(\d{2,4})(RS|S)?(?:\s*FP\s*2|FP2)?\b",
		text,
		flags=re.IGNORECASE,
	)
	if system_match:
		return f"{system_match.group(1)}{system_match.group(2) or ''}".upper()

	base = text.split()[0]
	if re.match(r"^K\d", base, flags=re.IGNORECASE):
		base = base[1:]

	return base.upper()


def contains_pattern(pattern, *values):
	text = " ".join(normalize_text(value) or "" for value in values)
	return bool(re.search(pattern, text, flags=re.IGNORECASE))


def word_match_variant(word_data):
	filename = word_data.get("Datei") or ""
	system = word_data.get("System")
	function = word_data.get("Funktion")

	if contains_pattern(r"FP\s*2", filename, system, function):
		return "FP2"

	if contains_pattern(r"SS04|SP[_\s.-]*0[.,]4|(?<!\d)0[.,]4(?!\d)", filename, system, function):
		return "SP_0.4"

	if contains_pattern(r"SS05|SP[_\s.-]*0[.,]5|(?<!\d)0[.,]5(?!\d)", filename, system, function):
		return "SP_0.5"

	if contains_pattern(r"\bVario\b|Change\s*Code|ChangeCode", filename, system, function):
		return "VARIO"

	if contains_pattern(r"OM\s*3|Omega\s*3", filename, system, function):
		return "OM3"

	if contains_pattern(r"OM\s*2|Omega\s*2", filename, system, function):
		return "OM2"

	if is_4000s_casa_word(system, filename):
		return "CASA"

	return "STANDARD"


def excel_match_variant(system_text, article_text):
	if contains_pattern(r"FP\s*2", system_text, article_text):
		return "FP2"

	if contains_pattern(r"Bolzenh[öo]he\s*0[.,]4|SP[_\s.-]*0[.,]4", system_text, article_text):
		return "SP_0.4"

	if contains_pattern(r"Bolzenh[öo]he\s*0[.,]5|SP[_\s.-]*0[.,]5|LV[_\s.-]*0[.,]5", system_text, article_text):
		return "SP_0.5"

	if contains_pattern(r"\bVario\b|Change\s*Code|ChangeCode", system_text, article_text):
		return "VARIO"

	if contains_pattern(r"OM\s*3|Omega\s*3", system_text, article_text):
		return "OM3"

	if contains_pattern(r"OM\s*2|Omega\s*2", system_text, article_text):
		return "OM2"

	if contains_pattern(r"\bCasa\b", system_text, article_text):
		return "CASA"

	return "STANDARD"


def normalize_diameter(value, field_name):
	return normalize_diameters(value, field_name)[0]


def normalize_diameters(value, field_name):
	text = normalize_text(value)
	if not text:
		raise ValueError(f"{field_name} fehlt")

	matches = re.findall(r"\d+(?:[.,]\d+)?", text)
	if not matches:
		raise ValueError(f"Cannot parse {field_name}: {value}")

	return list(dict.fromkeys(match.replace(",", ".") for match in matches))


def parse_rotor_diameter_and_type(value):
	text = normalize_text(value)
	if not text:
		raise ValueError("Rotor-diam fehlt")

	rotor_diameter = normalize_diameter(text, "Rotor-diam")
	rotor_text = text.lower()

	if re.search(r"exzentr|(?<![a-z])ex(?![a-z])", rotor_text):
		rotor_type = "ex"
	elif re.search(r"zentr|(?<![a-z])ze(?![a-z])", rotor_text):
		rotor_type = "ze"
	else:
		rotor_type = "ze"

	return rotor_diameter, rotor_type


def normalize_line(value):
	text = normalize_text(value)
	if not text:
		raise ValueError("Linie fehlt")

	match = re.search(r"\bL?\s*([0-9])\b", text, flags=re.IGNORECASE)
	if not match:
		raise ValueError(f"Cannot parse Linie: {value}")

	return match.group(1)


def system_line_override(word_data):
	filename = word_data.get("Datei") or ""

	if is_4000s_casa_word(word_data.get("System"), filename):
		return "4000S", "1"

	if is_4000s_grundausfuehrung_word(
		word_data.get("System"),
		filename,
		word_data.get("Funktion"),
	):
		return "4000S", "2"

	return None


def word_system_value(word_data):
	return word_data.get("System") or system_from_filename(word_data.get("Datei") or "")


def word_rotor_value(word_data):
	value = word_data.get("Rotor-diam")
	fallback = rotor_from_filename(word_data.get("Datei") or "")
	return fallback if fallback and not can_parse_rotor_value(value) else value


def word_hulse_value(word_data):
	value = word_data.get("Hulsen-diam")
	fallback = hulse_from_filename(word_data.get("Datei") or "")
	return fallback if fallback and not can_parse_diameter_value(value, "Hulsen-diam") else value


def parse_word_system_and_line(word_data):
	override = system_line_override(word_data)
	if override:
		return override

	filename = word_data.get("Datei") or ""
	system = normalize_system(word_system_value(word_data))
	try:
		line = normalize_line(word_data.get("Linie"))
	except ValueError:
		filename_line = line_from_filename(filename)
		if filename_line:
			return system, filename_line

		default_line = default_line_for_system(system)
		if default_line:
			return system, default_line[0]

		override = system_line_override(word_data)
		if override:
			return override

		raise

	return system, line


def parse_word_match_keys(word_data):
	rotor_diameter, rotor_type = parse_rotor_diameter_and_type(word_rotor_value(word_data))
	variant = word_match_variant(word_data)
	try:
		system, line = parse_word_system_and_line(word_data)
		lines = [line]
	except ValueError:
		if variant != "FP2":
			raise

		system = normalize_system(word_system_value(word_data))
		lines = ["1", "2"]
	hulse_diameters = normalize_diameters(word_hulse_value(word_data), "Hulsen-diam")

	return [
		(
			system,
			variant,
			rotor_diameter,
			rotor_type,
			hulse_diameter,
			line,
		)
		for hulse_diameter in hulse_diameters
		for line in lines
	]


def parse_word_match_key(word_data):
	return parse_word_match_keys(word_data)[0]


def parse_excel_match_keys(pair, include_se_matches=True):
	system_text, article_text = pair
	if not system_text or not article_text:
		return []

	article_match = re.search(
		r"^([A-Za-z]+)_([0-9]+(?:[.,][0-9]+)?)/([0-9]+(?:[.,][0-9]+)?)_(ze|ex)_(L[0-9](?:/L[0-9])?)(?:_|$)",
		str(article_text),
		flags=re.IGNORECASE,
	)
	if not article_match:
		return []

	prefix = article_match.group(1).upper()
	line_text = article_match.group(5).upper()

	if line_text == "L1/L2" and include_se_matches and prefix == "SE":
		lines = ["1", "2"]
	elif re.fullmatch(r"L[0-9]", line_text):
		lines = [line_text[1:]]
	else:
		lines = []

	return [
		(
			normalize_system(system_text),
			excel_match_variant(system_text, article_text),
			article_match.group(2).replace(",", "."),
			article_match.group(4).lower(),
			article_match.group(3).replace(",", "."),
			line,
		)
		for line in lines
	]


def parse_excel_match_key(pair):
	match_keys = parse_excel_match_keys(pair, include_se_matches=False)
	return match_keys[0] if match_keys else None


def find_matching_excel_rows(word_data, bolzenreihenset_pairs, include_se_matches=True):
	word_keys = parse_word_match_keys(word_data)
	return find_matching_excel_rows_by_keys(
		word_keys,
		bolzenreihenset_pairs,
		include_se_matches,
	)


def find_matching_excel_rows_by_key(word_key, bolzenreihenset_pairs, include_se_matches=True):
	return find_matching_excel_rows_by_keys(
		[word_key],
		bolzenreihenset_pairs,
		include_se_matches,
	)


def find_matching_excel_rows_by_keys(word_keys, bolzenreihenset_pairs, include_se_matches=True):
	word_key_set = set(word_keys)

	return [
		row_number
		for row_number, pair in bolzenreihenset_pairs.items()
		if word_key_set.intersection(parse_excel_match_keys(pair, include_se_matches))
	]


def build_unmatched_item(word_data, reason, search_key=None, matching_rows=None, error=None):
	return {
		"Grund": reason,
		"Referenz_zu_Abfullnorm": word_data.get("Referenz_zu_Abfullnorm"),
		"Suchschluessel": search_key,
		"Anzahl_Treffer": 0 if matching_rows is None else len(matching_rows),
		"Trefferzeilen": [] if matching_rows is None else matching_rows,
		"Fehler": error,
		"Word_Daten": word_data,
	}


def remove_conflicting_matches(matched, unmatched):
	matches_by_row = {}
	for item in matched:
		matches_by_row.setdefault(item["Excelzeile"], []).append(item)

	safe_matches = []
	for row_number, row_matches in matches_by_row.items():
		references = {item["Referenz_zu_Abfullnorm"] for item in row_matches}
		if len(references) == 1:
			safe_matches.extend(row_matches)
			continue

		conflict_details = [
			{
				"Referenz_zu_Abfullnorm": item["Referenz_zu_Abfullnorm"],
				"Word_Datei": item["Word_Daten"].get("Datei"),
				"Excelzeile": row_number,
				"Excel_Artikel": item["Excel_Artikel"],
			}
			for item in row_matches
		]

		for item in row_matches:
			unmatched.append(
				build_unmatched_item(
					item["Word_Daten"],
					"Excelzeile mehrdeutig: mehrere Referenzen",
					item.get("Suchschluessel"),
					[row_number],
					json.dumps(conflict_details, ensure_ascii=False),
				)
			)

	return safe_matches, unmatched


def match_references_to_excel_rows(
	word_data_list,
	bolzenreihenset_pairs,
	include_se_matches=True,
):
	matches = {}
	errors = []

	for word_data in word_data_list:
		matching_rows = find_matching_excel_rows(
			word_data,
			bolzenreihenset_pairs,
			include_se_matches,
		)
		reference = word_data["Referenz_zu_Abfullnorm"]

		if matching_rows:
			for row_number in matching_rows:
				matches[row_number] = reference
			continue

		errors.append(
			f"{reference}: expected at least 1 match, found {len(matching_rows)} "
			f"for key {parse_word_match_key(word_data)}; rows={matching_rows}"
		)

	if errors:
		raise ValueError("\n".join(errors))

	return matches


def build_match_report(
	word_data_list,
	bolzenreihenset_pairs,
	include_se_matches=True,
):
	matched = []
	unmatched = []

	for word_data in word_data_list:
		reference = word_data["Referenz_zu_Abfullnorm"]
		try:
			search_keys = parse_word_match_keys(word_data)
		except ValueError as exc:
			unmatched.append(
				build_unmatched_item(
					word_data,
					"Word-Datei nicht parsebar",
					error=str(exc),
				)
			)
			continue

		if not reference:
			unmatched.append(
				build_unmatched_item(
					word_data,
					"Referenz_zu_Abfullnorm fehlt",
					search_keys,
				)
			)
			continue

		matching_rows = find_matching_excel_rows_by_keys(
			search_keys,
			bolzenreihenset_pairs,
			include_se_matches,
		)

		if matching_rows:
			for row_number in matching_rows:
				excel_description, excel_article = bolzenreihenset_pairs[row_number]
				matched.append(
					{
						"Excelzeile": row_number,
						"Excel_Beschreibung": excel_description,
						"Excel_Artikel": excel_article,
						"Referenz_zu_Abfullnorm": reference,
						"Suchschluessel": search_keys,
						"Word_Daten": word_data,
					}
				)
			continue

		unmatched.append(
			build_unmatched_item(
				word_data,
				"kein passender Excel-Eintrag",
				search_keys,
				matching_rows,
			)
		)

	return remove_conflicting_matches(matched, unmatched)


def build_berechnungscode_match_items(
	word_data_list,
	bolzenreihenset_pairs,
	include_se_matches=True,
):
	items = []

	for word_data in word_data_list:
		try:
			search_keys = parse_word_match_keys(word_data)
		except ValueError:
			continue

		matching_rows = find_matching_excel_rows_by_keys(
			search_keys,
			bolzenreihenset_pairs,
			include_se_matches,
		)

		for row_number in matching_rows:
			excel_description, excel_article = bolzenreihenset_pairs[row_number]
			items.append(
				{
					"Excelzeile": row_number,
					"Excel_Beschreibung": excel_description,
					"Excel_Artikel": excel_article,
					"Referenz_zu_Abfullnorm": word_data.get("Referenz_zu_Abfullnorm"),
					"Suchschluessel": search_keys,
					"Word_Daten": word_data,
				}
			)

	return items


def write_unmatched_report(unmatched, path=UNMATCHED_FILE):
	lines = ["Nicht gematchte Referenzen", "==========================", ""]

	if not unmatched:
		lines.append("Alle auswertbaren Referenzen wurden gematcht.")
	else:
		counts = {}
		for item in unmatched:
			counts[item.get("Grund", "unbekannt")] = counts.get(item.get("Grund", "unbekannt"), 0) + 1

		lines.extend(["Zusammenfassung", "---------------"])
		for reason, count in sorted(counts.items()):
			lines.append(f"{reason}: {count}")
		lines.append("")

		for item in unmatched:
			lines.extend(
				[
					f"Grund: {item.get('Grund')}",
					f"Word-Datei: {item.get('Word_Daten', {}).get('Datei')}",
					f"Referenz: {item['Referenz_zu_Abfullnorm']}",
					f"Suchschluessel: {item['Suchschluessel']}",
					f"Anzahl Treffer: {item['Anzahl_Treffer']}",
					f"Trefferzeilen: {item['Trefferzeilen']}",
					f"Fehler: {item.get('Fehler')}",
					"Word_Daten:",
					json.dumps(item["Word_Daten"], indent=2, ensure_ascii=False),
					"",
				]
			)

	path.write_text("\n".join(lines), encoding="utf-8")
	return path


def clean_table_cell(value):
	return value.replace("\r", " ").replace("\t", " ").strip()


def is_calculation_code(value):
	return bool(re.fullmatch(r"[0-9A-Z]", value))


def drawing_numbers_from_cell(value):
	return re.findall(r"(?:15[012]|54)\.\d+(?:\.[0-9A-Za-z]+)*", value)


def first_drawing_with_prefix(drawings, prefix, start_index=0):
	for index in range(start_index, len(drawings)):
		if drawings[index].startswith(prefix):
			return index, drawings[index]

	return None, None


def extract_word_table_section(text, start_marker, end_marker):
	start = text.find(start_marker)
	if start < 0:
		return ""

	end = text.find(end_marker, start + len(start_marker))
	if end < 0:
		return text[start:]

	return text[start:end]


def parse_fill_table_rows(section_text):
	cells = [clean_table_cell(cell) for cell in section_text.split("\x07")]
	cells = [cell for cell in cells if cell]
	rows = []
	index = 0

	while index < len(cells) - 1:
		if not is_calculation_code(cells[index]) or not drawing_numbers_from_cell(cells[index + 1]):
			index += 1
			continue

		code = cells[index]
		next_index = index + 2
		while next_index < len(cells):
			if is_calculation_code(cells[next_index]) and next_index + 1 < len(cells):
				if drawing_numbers_from_cell(cells[next_index + 1]):
					break
			next_index += 1

		row_cells = cells[index:next_index]
		drawings = []
		for cell in row_cells[1:]:
			drawings.extend(drawing_numbers_from_cell(cell))

		bolzen_index, bolzen_drawing = first_drawing_with_prefix(drawings, "152")
		if bolzen_drawing is None:
			bolzen_index, bolzen_drawing = first_drawing_with_prefix(drawings, "54")

		gegenbolzen_index, gegenbolzen_drawing = first_drawing_with_prefix(
			drawings,
			"151",
			0 if bolzen_index is None else bolzen_index + 1,
		)
		extra_drawings = [
			drawing
			for drawing_index, drawing in enumerate(drawings)
			if drawing_index not in {bolzen_index, gegenbolzen_index}
		]

		rows.append(
			{
				"Code": code,
				"Bolzen_Zeichnungsnummer": bolzen_drawing,
				"Gegenbolzen_Zeichnungsnummer": gegenbolzen_drawing,
				"Weitere_Zeichnungsnummern": extra_drawings,
				"Rohzellen": row_cells,
			}
		)
		index = next_index

	return rows


def parse_fill_tables(path):
	text = read_word_text(path)
	seiten_text = extract_word_table_section(text, "Seitenabf", "Hochkantabf")
	hochkant_text = extract_word_table_section(text, "Hochkantabf", "Hinweis:")

	return {
		"SE": parse_fill_table_rows(seiten_text),
		"HK": parse_fill_table_rows(hochkant_text),
	}


def load_drawing_lookup(workbook):
	sheet = workbook[DRAWING_LOOKUP_SHEET_NAME]
	lookup = {}

	for drawing_article, drawing_number, *_ in sheet.iter_rows(min_row=2, values_only=True):
		if drawing_article and drawing_number:
			lookup[str(drawing_number).strip()] = str(drawing_article).strip()

	return lookup


def article_prefix(article):
	match = re.match(r"^([A-Za-z]+)_", article or "")
	return match.group(1).upper() if match else None


def build_word_path_by_reference(word_data_list, folder=DOCX_FOLDER):
	return {
		word_data["Referenz_zu_Abfullnorm"]: folder / word_data["Datei"]
		for word_data in word_data_list
		if word_data.get("Referenz_zu_Abfullnorm") and word_data.get("Datei")
	}


def production_system_from_description(description):
	if not description:
		return None

	return str(description).split()[0]


def production_systems_by_bolzenreihenset(bolzenreihenset_pairs):
	systems_by_bolzenreihenset = {}

	for description, bolzenreihenset in bolzenreihenset_pairs.values():
		production_system = production_system_from_description(description)
		if not production_system or not bolzenreihenset:
			continue

		systems_by_bolzenreihenset.setdefault(str(bolzenreihenset).strip(), set()).add(production_system)

	return systems_by_bolzenreihenset


def calculation_rows_by_article_and_code(sheet):
	rows = {}

	for row_number in range(2, sheet.max_row + 1):
		article = sheet.cell(row_number, 2).value
		code = sheet.cell(row_number, 3).value
		if article is None or code is None:
			continue

		rows.setdefault((str(article).strip(), str(code).strip()), []).append(row_number)

	return rows


def find_calculation_rows(
	calc_rows,
	excel_article,
	code,
	production_system,
	systems_by_bolzenreihenset,
):
	row_numbers = calc_rows.get((excel_article, code), [])
	known_systems = systems_by_bolzenreihenset.get(excel_article, set())

	if production_system not in known_systems:
		return [], f"Produktionssystem passt nicht zum Bolzenreihenset: {production_system} nicht in {sorted(known_systems)}"

	return row_numbers, None


def first_calculation_row_for_article(calc_rows, excel_article):
	for (article, _code), row_numbers in calc_rows.items():
		if article == excel_article and row_numbers:
			return row_numbers[-1]

	return None


def copy_row_layout(sheet, source_row_number, target_row_number):
	if source_row_number is None:
		return

	sheet.row_dimensions[target_row_number].height = sheet.row_dimensions[source_row_number].height
	for column_number in range(1, max(sheet.max_column, 13) + 1):
		source_cell = sheet.cell(source_row_number, column_number)
		target_cell = sheet.cell(target_row_number, column_number)
		if source_cell.has_style:
			target_cell._style = copy(source_cell._style)
		target_cell.number_format = source_cell.number_format
		target_cell.protection = copy(source_cell.protection)
		target_cell.alignment = copy(source_cell.alignment)


def production_system_formula(row_number):
	return f"=_xlfn.XLOOKUP(B{row_number},'[1]Artikelstamm bis Bolzen'!$F:$F,'[1]Artikelstamm bis Bolzen'!$C:$C,\"\")"


def expand_tables_to_row(sheet, row_number):
	for table_name in list(sheet.tables.keys()):
		table = sheet.tables[table_name]
		min_column, min_row, max_column, max_row = range_boundaries(table.ref)
		if row_number <= max_row:
			continue

		table.ref = f"{get_column_letter(min_column)}{min_row}:{get_column_letter(max_column)}{row_number}"
		if table.autoFilter:
			table.autoFilter.ref = table.ref


def append_calculation_code_row(sheet, calc_rows, excel_article, code):
	new_row_number = sheet.max_row + 1
	template_row_number = first_calculation_row_for_article(calc_rows, excel_article) or sheet.max_row
	copy_row_layout(sheet, template_row_number, new_row_number)

	sheet.cell(new_row_number, 1).value = production_system_formula(new_row_number)
	sheet.cell(new_row_number, 2).value = excel_article
	sheet.cell(new_row_number, 3).value = code
	for column_number in range(4, 14):
		sheet.cell(new_row_number, column_number).value = None

	expand_tables_to_row(sheet, new_row_number)
	calc_rows.setdefault((excel_article, code), []).append(new_row_number)
	return new_row_number


def map_fill_row_to_articles(fill_row, drawing_lookup):
	reasons = []
	bolzen_drawing = fill_row["Bolzen_Zeichnungsnummer"]
	gegenbolzen_drawing = fill_row["Gegenbolzen_Zeichnungsnummer"]
	bolzen_article = drawing_lookup.get(bolzen_drawing) if bolzen_drawing else None
	gegenbolzen_article = drawing_lookup.get(gegenbolzen_drawing) if gegenbolzen_drawing else None

	if not bolzen_drawing:
		reasons.append("keine Bolzen-Zeichnungsnummer erkannt")
	elif not bolzen_article:
		reasons.append(f"Bolzen-Zeichnungsnummer nicht im Lookup: {bolzen_drawing}")

	if not gegenbolzen_drawing:
		reasons.append("keine Gegenbolzen-Zeichnungsnummer erkannt")
	elif not gegenbolzen_article:
		reasons.append(f"Gegenbolzen-Zeichnungsnummer nicht im Lookup: {gegenbolzen_drawing}")

	return bolzen_article, gegenbolzen_article, reasons


def target_pair_columns(pair_index):
	start_column = 4 + pair_index * 2
	return start_column, start_column + 1


def is_empty_cell_value(value):
	return value in {None, ""}


def write_pair_to_next_available_slot(sheet, row_number, bolzen_article, gegenbolzen_article, context, written, report_items):
	for pair_index in range(5):
		bolzen_column, gegenbolzen_column = target_pair_columns(pair_index)
		bolzen_cell = sheet.cell(row_number, bolzen_column)
		gegenbolzen_cell = sheet.cell(row_number, gegenbolzen_column)

		if bolzen_cell.value == bolzen_article and gegenbolzen_cell.value == gegenbolzen_article:
			return False

		if is_empty_cell_value(bolzen_cell.value) and is_empty_cell_value(gegenbolzen_cell.value):
			bolzen_cell.value = bolzen_article
			gegenbolzen_cell.value = gegenbolzen_article
			written.append(
				{
					**context,
					"Bolzen_Artikel": bolzen_article,
					"Gegenbolzen_Artikel": gegenbolzen_article,
					"Bolzen_Zelle": bolzen_cell.coordinate,
					"Gegenbolzen_Zelle": gegenbolzen_cell.coordinate,
				}
			)
			return True

	report_items.append(
		{
			**context,
			"Grund": "kein freies Bolzen/Gegenbolzen-Paar in D-M",
			"Bolzen_Artikel": bolzen_article,
			"Gegenbolzen_Artikel": gegenbolzen_article,
		}
	)
	return False


def write_or_report_cell(sheet, row_number, column_number, value, context, report_items):
	existing_value = sheet.cell(row_number, column_number).value
	if existing_value in {None, "", value}:
		sheet.cell(row_number, column_number).value = value
		return True

	report_items.append(
		{
			**context,
			"Grund": "Zielzelle enthält abweichenden Wert",
			"Zelle": sheet.cell(row_number, column_number).coordinate,
			"Vorhanden": existing_value,
			"Neu": value,
		}
	)
	return False


def write_berechnungscode_report(report_items, path=CALC_REPORT_FILE):
	lines = ["Nicht geschriebene Berechnungscode-Eintraege", "===========================================", ""]

	if not report_items:
		lines.append("Alle sicheren Berechnungscode-Eintraege wurden geschrieben.")
	else:
		for item in report_items:
			lines.extend(
				[
					f"Grund: {item.get('Grund')}",
					f"Referenz: {item.get('Referenz_zu_Abfullnorm')}",
					f"Word-Datei: {item.get('Word_Datei')}",
					f"Excel-Artikel: {item.get('Excel_Artikel')}",
					f"Tabelle: {item.get('Tabelle')}",
					f"Code: {item.get('Code')}",
					f"Excelzeile Berechnungscode: {item.get('Berechnungscode_Zeile')}",
					f"Details: {json.dumps(item, ensure_ascii=False)}",
					"",
				]
			)

	path.write_text("\n".join(lines), encoding="utf-8")
	return path


def write_berechnungscode_entries(
	matched,
	word_data_list,
	path=XLSX_FILE,
	report_path=CALC_REPORT_FILE,
	bolzenreihenset_pairs=None,
):
	if bolzenreihenset_pairs is None:
		bolzenreihenset_pairs = load_bolzenreihenset_pairs(path)

	word_path_by_reference = build_word_path_by_reference(word_data_list)
	systems_by_bolzenreihenset = production_systems_by_bolzenreihenset(bolzenreihenset_pairs)
	fill_tables_by_path = {}
	written = []
	report_items = []
	workbook = load_workbook(path)
	calc_sheet = workbook[CALC_SHEET_NAME]
	drawing_lookup = load_drawing_lookup(workbook)
	calc_rows = calculation_rows_by_article_and_code(calc_sheet)

	try:
		for match_item in matched:
			reference = match_item["Referenz_zu_Abfullnorm"]
			excel_article = match_item["Excel_Artikel"]
			production_system = production_system_from_description(match_item["Excel_Beschreibung"])
			prefix = article_prefix(excel_article)
			word_filename = match_item.get("Word_Daten", {}).get("Datei")
			word_path = DOCX_FOLDER / word_filename if word_filename else word_path_by_reference.get(reference)

			base_context = {
				"Referenz_zu_Abfullnorm": reference,
				"Word_Datei": None if word_path is None else word_path.name,
				"Excel_Artikel": excel_article,
				"Produktionssystem": production_system,
				"Tabelle": prefix,
			}

			if prefix not in {"HK", "SE"}:
				report_items.append({**base_context, "Grund": "Excel-Artikel beginnt nicht mit HK oder SE"})
				continue

			if word_path is None or not word_path.exists():
				report_items.append({**base_context, "Grund": "zugehoerige Word-Datei nicht gefunden"})
				continue

			if word_path not in fill_tables_by_path:
				fill_tables_by_path[word_path] = parse_fill_tables(word_path)

			rows_by_code = {}
			for fill_row in fill_tables_by_path[word_path][prefix]:
				rows_by_code.setdefault(fill_row["Code"], []).append(fill_row)

			for code, fill_rows in rows_by_code.items():
				calculation_row_numbers, system_error = find_calculation_rows(
					calc_rows,
					excel_article,
					code,
					production_system,
					systems_by_bolzenreihenset,
				)
				if system_error:
					report_items.append(
						{
							**base_context,
							"Code": code,
							"Grund": system_error,
							"Trefferzeilen": calculation_row_numbers,
						}
					)
					continue

				if not calculation_row_numbers:
					calculation_row_number = append_calculation_code_row(
						calc_sheet,
						calc_rows,
						excel_article,
						code,
					)
				elif len(calculation_row_numbers) != 1:
					report_items.append(
						{
							**base_context,
							"Code": code,
							"Grund": f"Berechnungscode-Zeile nicht eindeutig gefunden: {len(calculation_row_numbers)} Treffer",
							"Trefferzeilen": calculation_row_numbers,
						}
					)
					continue
				else:
					calculation_row_number = calculation_row_numbers[0]
				for fill_row in fill_rows:
					context = {
						**base_context,
						"Code": code,
						"Berechnungscode_Zeile": calculation_row_number,
						"Word_Tabellenzeile": fill_row,
					}
					bolzen_article, gegenbolzen_article, reasons = map_fill_row_to_articles(fill_row, drawing_lookup)

					if reasons:
						report_items.append({**context, "Grund": "; ".join(reasons)})
						continue

					write_pair_to_next_available_slot(
						calc_sheet,
						calculation_row_number,
						bolzen_article,
						gegenbolzen_article,
						context,
						written,
						report_items,
					)

		workbook.save(path)
	finally:
		workbook.close()

	write_berechnungscode_report(report_items, report_path)
	return written, report_items


def write_references_to_excel(
	path=XLSX_FILE,
	unmatched_path=UNMATCHED_FILE,
	include_se_matches=True,
	word_data_list=None,
	bolzenreihenset_pairs=None,
):
	if word_data_list is None:
		word_data_list = load_all_word_data()
	if bolzenreihenset_pairs is None:
		bolzenreihenset_pairs = load_bolzenreihenset_pairs(path)

	matched, unmatched = build_match_report(
		word_data_list,
		bolzenreihenset_pairs,
		include_se_matches,
	)

	workbook = load_workbook(path)
	sheet = workbook[ABFUELLNORM_SHEET_NAME]

	try:
		for item in matched:
			row_number = item["Excelzeile"]
			reference = item["Referenz_zu_Abfullnorm"]
			sheet[f"C{row_number}"] = reference

		workbook.save(path)
	finally:
		workbook.close()

	write_unmatched_report(unmatched, unmatched_path)
	return matched, unmatched


if __name__ == "__main__":
	word_data = load_all_word_data()
	bolzenreihenset_pairs = load_bolzenreihenset_pairs()

	matched_data, unmatched_data = write_references_to_excel(
		word_data_list=word_data,
		bolzenreihenset_pairs=bolzenreihenset_pairs,
	)
	berechnungscode_matches = build_berechnungscode_match_items(
		word_data,
		bolzenreihenset_pairs,
	)
	berechnungscode_data, berechnungscode_report = write_berechnungscode_entries(
		berechnungscode_matches,
		word_data,
		bolzenreihenset_pairs=bolzenreihenset_pairs,
	)

	print(f"{len(word_data)} Word-Dateien gelesen")
	print(f"{len(matched_data)} sichere Referenz-Eintraege in Excel geschrieben")
	print(f"{len(berechnungscode_data)} sichere Berechnungscode-Eintraege in Excel geschrieben")
	if unmatched_data:
		print(
			f"{len(unmatched_data)} nicht geschriebene Referenz-Eintraege stehen in {UNMATCHED_FILE}",
			file=sys.stderr,
		)
	if berechnungscode_report:
		print(
			f"{len(berechnungscode_report)} nicht geschriebene Berechnungscode-Eintraege stehen in {CALC_REPORT_FILE}",
			file=sys.stderr,
		)